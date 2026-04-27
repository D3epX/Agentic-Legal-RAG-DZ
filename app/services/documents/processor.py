"""
Document processor — text extraction and intelligent chunking.

Supports PDF, TXT, DOCX and XLSX.  Chunks are sized to fit within
embedding model context windows while preserving sentence boundaries.
"""

import io
import logging
import re
from typing import List, Dict, Optional

import PyPDF2

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Extract text from files and split into overlapping chunks."""

    def __init__(
        self,
        chunk_size: int = settings.CHUNK_SIZE,
        chunk_overlap: int = settings.CHUNK_OVERLAP,
        max_chunks: int = settings.MAX_CHUNKS_PER_DOC,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chunks = max_chunks

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    def extract_pdf_text(self, pdf_bytes: bytes) -> List[Dict]:
        """Extract text from PDF, returning list of {page, content}."""
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        pages: List[Dict] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = self._normalise(text)
            if text.strip():
                pages.append({"page": idx, "content": text})
        return pages

    def extract_docx_text(self, docx_bytes: bytes) -> List[Dict]:
        """Extract text from a DOCX file, returning list of {page, content}.

        Each paragraph becomes part of a single logical 'page' because
        DOCX files don't have physical page breaks we can reliably detect.
        We split on every ~3000 chars to create manageable chunks.
        """
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    full_text += "\n" + row_text

        full_text = self._normalise(full_text)
        if not full_text.strip():
            return []

        # Split into virtual pages (~3000 chars each)
        pages: List[Dict] = []
        chunk_size = 3000
        for i in range(0, len(full_text), chunk_size):
            segment = full_text[i : i + chunk_size].strip()
            if segment:
                pages.append({"page": len(pages) + 1, "content": segment})
        return pages

    def extract_xlsx_text(self, xlsx_bytes: bytes) -> List[Dict]:
        """Extract text from an XLSX file, one 'page' per sheet."""
        import openpyxl

        wb = openpyxl.load_workbook(
            io.BytesIO(xlsx_bytes), read_only=True, data_only=True
        )
        pages: List[Dict] = []
        for idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]
            rows_text: List[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(c).strip() for c in row if c is not None and str(c).strip()
                ]
                if cells:
                    rows_text.append("\t".join(cells))
            if rows_text:
                content = f"Sheet: {sheet_name}\n" + "\n".join(rows_text)
                pages.append({"page": idx, "content": self._normalise(content)})
        wb.close()
        return pages

    def extract_text(self, raw_text: str) -> List[Dict]:
        """Wrap raw text string as a single page."""
        return [{"page": 1, "content": self._normalise(raw_text)}]

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def chunk_text(
        self,
        text: str,
        page_map: Optional[List[Dict]] = None,
    ) -> List[Dict]:
        """Split text into overlapping chunks.

        Returns list of {"content": str, "page": int | None}.
        """
        if page_map:
            return self._chunk_pages(page_map)
        return self._chunk_flat(text)

    def chunk_pdf(self, pdf_bytes: bytes) -> List[Dict]:
        pages = self.extract_pdf_text(pdf_bytes)
        return self._chunk_pages(pages)

    # ------------------------------------------------------------------
    # Phase 7: Legal-aware chunking
    # ------------------------------------------------------------------

    # Regex for legal article headings (multilingual)
    _LEGAL_HEADING_RE = re.compile(
        r"(?:^|\n)\s*(?:"
        r"(?:Art(?:icle)?\.?\s*\d[\d\w.\-/]*)"  # Article / Art. 43, Art.3-1
        r"|(?:المادة\s*\d[\d\w.\-/]*)"            # Arabic: المادة 43
        r"|(?:Section\s+\d[\d.]*)"                 # Section 3.2
        r"|(?:Chapitre\s+\w+)"                     # Chapitre III
        r"|(?:الفصل\s+\w+)"                        # Arabic: chapter
        r"|(?:Titre\s+\w+)"                        # Titre II
        r"|(?:الباب\s+\w+)"                         # Arabic: title/part
        r")",
        re.IGNORECASE | re.UNICODE,
    )

    def chunk_legal(
        self,
        text: str,
        page_map: Optional[List[Dict]] = None,
    ) -> List[Dict]:
        """Structure-aware chunking for legal documents.

        Priority order:
          1. Split at legal article headings (Article X, المادة, etc.)
          2. If articles are too large, split at paragraph boundaries
          3. If paragraphs are too large, fall back to sentence-aware split

        Each chunk includes metadata: article heading if detected.
        """
        # Combine all pages into a single text if page_map is provided
        if page_map:
            full_text = "\n\n".join(p["content"] for p in page_map)
        else:
            full_text = text

        # Try to split on article boundaries
        articles = self._split_on_articles(full_text)

        if len(articles) < 2:
            # No article structure detected → fall back to standard chunking
            logger.debug("No article structure found, using standard chunking")
            if page_map:
                return self._chunk_pages(page_map)
            return self._chunk_flat(full_text)

        # Process each article
        chunks: List[Dict] = []
        for article in articles:
            heading = article.get("heading")
            content = article["content"].strip()

            if not content or len(content.split()) < 10:
                continue

            # If article fits within chunk_size, keep it as one chunk
            word_count = len(content.split())
            if word_count <= self.chunk_size:
                chunks.append({
                    "content": content,
                    "page": None,
                    "article_heading": heading,
                })
            else:
                # Article too large — split within it
                sub_chunks = self._split(content)
                for i, sub in enumerate(sub_chunks):
                    chunk_heading = heading
                    if i > 0 and heading:
                        chunk_heading = f"{heading} (cont.)"
                    chunks.append({
                        "content": sub,
                        "page": None,
                        "article_heading": chunk_heading,
                    })

            if len(chunks) >= self.max_chunks:
                break

        logger.info(
            "Legal chunking: %d articles → %d chunks",
            len(articles), len(chunks),
        )
        return chunks

    def _split_on_articles(self, text: str) -> List[Dict]:
        """Split text at legal article boundaries.

        Returns a list of {"heading": str|None, "content": str}.
        """
        matches = list(self._LEGAL_HEADING_RE.finditer(text))

        if not matches:
            return [{"heading": None, "content": text}]

        articles: List[Dict] = []

        # Content before first article heading (preamble/intro)
        if matches[0].start() > 0:
            preamble = text[:matches[0].start()].strip()
            if preamble and len(preamble.split()) > 10:
                articles.append({"heading": "Preamble", "content": preamble})

        # Each article section
        for i, match in enumerate(matches):
            heading = match.group().strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()

            # Include the heading in the content for context
            full_content = f"{heading}\n{content}"
            articles.append({"heading": heading, "content": full_content})

        return articles


    # ------------------------------------------------------------------
    # Internal
    # ------------------------------------------------------------------

    def _chunk_pages(self, pages: List[Dict]) -> List[Dict]:
        chunks: List[Dict] = []
        for page_info in pages:
            page_chunks = self._split(page_info["content"])
            for c in page_chunks:
                chunks.append({"content": c, "page": page_info["page"]})
                if len(chunks) >= self.max_chunks:
                    return chunks
        return chunks

    def _chunk_flat(self, text: str) -> List[Dict]:
        parts = self._split(text)
        return [{"content": p, "page": None} for p in parts[: self.max_chunks]]

    def _split(self, text: str) -> List[str]:
        """Sliding-window sentence-aware splitter.

        Tries to break on sentence boundaries ('.', '!', '?', Arabic period)
        then falls back to word boundaries.
        """
        text = text.strip()
        if not text:
            return []

        sentences = re.split(r"(?<=[.!?。؟])\s+", text)
        chunks: List[str] = []
        current: List[str] = []
        current_len = 0

        for sent in sentences:
            sent_len = len(sent.split())
            if current_len + sent_len > self.chunk_size and current:
                chunks.append(" ".join(current))
                # Overlap: keep tail sentences
                overlap_words = 0
                overlap_start = len(current)
                for i in range(len(current) - 1, -1, -1):
                    overlap_words += len(current[i].split())
                    if overlap_words >= self.chunk_overlap:
                        overlap_start = i
                        break
                current = current[overlap_start:]
                current_len = sum(len(s.split()) for s in current)

            current.append(sent)
            current_len += sent_len

        if current:
            chunks.append(" ".join(current))

        return chunks

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean raw text before chunking.

        Phase 7: centralised cleaning step applied to all user-uploaded
        documents (both PDF-extracted and plain-text uploads).
        """
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _normalise(text: str) -> str:
        """Basic whitespace normalisation."""
        return re.sub(r"\s+", " ", text).strip()


# Singleton
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
